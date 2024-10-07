import tkinter as tk
import os
import sys
import logging
import datetime
import threading
from tkinter import filedialog, messagebox
from src.back.calculations import Calculations
from src.back.file_handler import FileHandler
from src.front.utils import Utils
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
import PyPDF2
import customtkinter as ctk  # Asegúrate de importar customtkinter

if os.path.exists("error_log.txt"):
    os.remove(path="error_log.txt")

logging.basicConfig(
    filename="error_log.txt",
    filemode="a",
    format="[%(asctime)s] - [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    level=logging.ERROR,
)


class GuiServices:
    def __init__(self, frame_main) -> None:
        self.window = frame_main
        self.sensor_number = 0
        self.sensor_data = {}
        self.utils = Utils(self.window)
        self.calculations = Calculations(self)
        self.operations_steps: str
        self.files_frame = None

    def set_files_frame(self, frame):
        """Establecer el marco donde se muestran los archivos cargados."""
        self.files_frame = frame

    def load_files(self):
        filepaths = filedialog.askopenfilenames(
            filetypes=[
                ("All files", "*"),
                ("CSV files", "*.csv"),
                ("TXT files", "*.txt"),
            ]
        )

        if filepaths:
            try:
                f = FileHandler(self)
                self.sensor_data = f.load_csv_files(list(filepaths))
                self.sensor_number = f.count_sensors()

                self.update_files_ui(filepaths)

                return self.sensor_data

            except Exception as e:
                return None

        else:
            return None

    def update_files_ui(self, filepaths):
        """Actualizar la UI para ocultar el botón de selección de archivos y mostrar los nombres."""
        if self.files_frame:
            for widget in self.files_frame.winfo_children():
                widget.destroy()

            for filepath in filepaths:
                filename = os.path.basename(filepath)
                label = tk.Label(self.files_frame, text=filename, bg="white")
                label.pack(pady=5)

    def drop_frame(self, frame_main):
        for widget in frame_main.winfo_children():
            if isinstance(widget, tk.Frame):
                self.drop_frame(widget)
            widget.destroy()

    def get_values(self, frame):
        values = []
        for widget in frame.winfo_children():
            if isinstance(widget, tk.Entry):
                new_values = widget.get().replace(",", ".")
                values.append(new_values)
            elif isinstance(widget, tk.Frame):
                values.extend(self.get_values(widget))

        return values

    def create_result(self, distances):
        self.results = self.calculations.calculate_magnetic_moment(
            self.sensor_data, distances
        )
        return self.results

    def get_list_paths(self):
        self.paths = self.calculations.get_image_paths()
        return self.paths

    def obtain_calculation_details(self):
        details = self.calculations.get_calculation_steps()
        return details

    def centre_text_cell_centre(self, cell, text):
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)

    def modify_word_table(self, filepath, nueva_ruta):
        try:
            doc = Document(filepath)
            table = doc.tables[1]

            self.centre_text_cell_centre(table.cell(1, 1), "9.8 ± 4.2")
            self.centre_text_cell_centre(table.cell(1, 2), "5.7 ± 1.5")
            self.centre_text_cell_centre(table.cell(1, 3), "5.7 ± 1.5")
            self.centre_text_cell_centre(table.cell(1, 4), "5.7 ± 1.5")
            self.centre_text_cell_centre(table.cell(1, 5), "< 60")

            self.centre_text_cell_centre(table.cell(2, 1), "9.8 ± 4.2")
            self.centre_text_cell_centre(table.cell(2, 2), "5.7 ± 1.5")
            self.centre_text_cell_centre(table.cell(2, 3), "5.7 ± 1.5")
            self.centre_text_cell_centre(table.cell(2, 4), "5.7 ± 1.5")
            self.centre_text_cell_centre(table.cell(2, 5), "< 55")

            self.centre_text_cell_centre(table.cell(3, 1), "9.8 ± 4.2")
            self.centre_text_cell_centre(table.cell(3, 2), "5.7 ± 1.5")
            self.centre_text_cell_centre(table.cell(3, 3), "5.7 ± 1.5")
            self.centre_text_cell_centre(table.cell(3, 4), "5.7 ± 1.5")
            self.centre_text_cell_centre(table.cell(3, 5), "< 45")

            doc.save(nueva_ruta)
            return True

        except Exception as e:
            print("Error in modify_word_table", str(e))
            return False

    def convert_word_to_pdf(self, path_word, pdf_path):
        try:
            convert(path_word, pdf_path)
            return True
        except Exception as e:
            self.log_error("Error", f"Error in convert_word_to_pdf {e}")
            return False

    def modify_pdf_metadata(self, pdf_path):
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            writer = PyPDF2.PdfWriter()

            for page in reader.pages:
                writer.add_page(page)

            metadatos = reader.metadata
            metadatos.update({
                '/Title': pdf_path
            })

            writer.add_metadata(metadatos)

            with open(f"{pdf_path}", 'wb') as output_file:
                writer.write(output_file)
        
        return f"{pdf_path}"

    def async_convert_and_modify_pdf(self, new_path_word, pdf_path, progress_bar, progress_window):
        progress_bar.set(0.5)  # Actualiza la barra de progreso al 50% cuando comienza la conversión

        if self.convert_word_to_pdf(new_path_word, pdf_path):
            progress_bar.set(0.75)  # Actualiza la barra de progreso al 75%
            modified_pdf = self.modify_pdf_metadata(pdf_path)
            if modified_pdf:
                os.startfile(modified_pdf)  # Abrir el PDF resultante
            os.remove(new_path_word)  # Eliminar el archivo Word temporal

        progress_bar.set(1)  # Progreso completo
        progress_window.destroy()  # Cerrar la ventana de progreso

    def export_to_pdf(self):
        word_file = self.resource_path("resource\\Ejemplo_Informe_de_resultados.docx")
        if word_file:
            new_path_word = self.resource_path("resource\\tabla_modificada.docx")
            now = datetime.datetime.now()
            fecha_hora = now.strftime("%Y-%m-%d_%H-%M-%S")
            pdf_path = f"Results_report_{fecha_hora}.pdf"

            if self.modify_word_table(word_file, new_path_word):
                # Crear ventana de progreso
                progress_window = ctk.CTkToplevel(self.window)
                progress_window.title("Generando PDF...")
                progress_window.geometry("300x100")

                progress_label = ctk.CTkLabel(progress_window, text="Generando el PDF, por favor espere...")
                progress_label.pack(pady=10)

                progress_bar = ctk.CTkProgressBar(progress_window)
                progress_bar.pack(pady=20, padx=10)
                progress_bar.set(0)

                # Crear un hilo para la conversión a PDF y modificación de metadatos
                threading.Thread(target=self.async_convert_and_modify_pdf, args=(new_path_word, pdf_path, progress_bar, progress_window)).start()
            else:
                self.show_message("Could not change the document table")

    def show_message(self, msg):
        messagebox.showerror("Error", msg)

    def log_error(self, error_type, error_message):
        logger.error(f"{error_type}: {error_message}")
        self.show_message(f"{error_type}: {error_message}")

    def resource_path(self, relative_path):
        try:
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")

        return os.path.join(base_path, relative_path)
