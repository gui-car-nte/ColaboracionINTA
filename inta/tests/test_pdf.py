from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
# from docx2pdf import convert


def centrar_texto_celda(celda, texto):
    """Centrar el texto horizontal y verticalmente dentro de una celda."""
    celda.text = texto
    # Centrar horizontalmente
    for paragraph in celda.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Centrar verticalmente
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def modificar_tabla_word(filepath, nueva_ruta):
    """Modificar una tabla en un archivo Word y guardar los cambios."""
    try:
        # Abrir el documento
        doc = Document(filepath)
        # Seleccionar la segunda tabla (índice 1)
        table = doc.tables[1]

        # Modificar el contenido de la tabla y centrar el texto
        centrar_texto_celda(table.cell(1, 1), '9.8 ± 4.2')
        centrar_texto_celda(table.cell(1, 2), '5.7 ± 1.5')
        centrar_texto_celda(table.cell(1, 3), '5.7 ± 1.5')
        centrar_texto_celda(table.cell(1, 4), '5.7 ± 1.5')
        centrar_texto_celda(table.cell(1, 5), '< 60')

        # Guardar el documento modificado
        doc.save(nueva_ruta)
        return True

    except Exception as e:
        print("Error en modificar_tabla_word", str(e))
        return False

def convertir_word_a_pdf(ruta_word, ruta_pdf):
    """Convertir un archivo .docx a PDF."""
    try:
        from docx2pdf import convert
        convert(ruta_word, ruta_pdf)
        return True
    except Exception as e:
        print("Error en convertir_word_a_pdf", str(e))
        return False

def export_to_pdf():
    """Ejecutar el proceso de modificar una tabla en un Word y convertirlo a PDF."""
    archivo_word = "src/front/resource/Ejemplo_Informe_de_resultados.docx"
    if archivo_word:
        nueva_ruta_word = "tabla_modificada.docx"
        ruta_pdf = "tabla_modificada.pdf"
        
        if modificar_tabla_word(archivo_word, nueva_ruta_word):
            if convertir_word_a_pdf(nueva_ruta_word, ruta_pdf):
                os.startfile(ruta_pdf)  # Abrir el PDF resultante
                os.remove(nueva_ruta_word)  # Eliminar el archivo Word temporal
            else:
                print("No se pudo convertir el documento a PDF", "red")
        else:
            print("No se pudo modificar la tabla del documento", "red")


export_to_pdf()